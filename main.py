from docx import Document
from fuzzywuzzy import fuzz
from docx.shared import RGBColor
from docx.enum.text import WD_COLOR

# Загрузить документы
doc_text = Document('text1.docx')
doc_dict = Document('Dictionary1.docx')

# Выставление порога схожести
similarity_threshold = 60

dictionary_phrases = [p.text for p in doc_dict.paragraphs]


def highlight_phrase(paragraph, phrase, color_rgb, count):
    """
    Функция для выделения фразы в параграфе заданным цветом.
    """
    if phrase in paragraph.text:
        # Разделить параграф по тексту, который нужно выделить
        pre, match, post = paragraph.text.partition(phrase)

        # Очистить параграф и добавить в него части с выделением
        paragraph.clear()
        paragraph.add_run(pre)
        highlighted = paragraph.add_run(match +f'  [{count}]')
        highlighted.font.highlight_color = color_rgb
        paragraph.add_run(post)


# Перебрать все параграфы в документе и проверить на совпадение с фразами из словаря
for paragraph in doc_text.paragraphs:
    for count, phrase in enumerate(dictionary_phrases):
        # В данном примере используем fuzzy matching с порогом в similarity_threshold%
        if fuzz.partial_ratio(paragraph.text, phrase) >= similarity_threshold:
            highlight_phrase(paragraph, phrase, WD_COLOR.YELLOW, count)  # желтый цвет

    # paragraph.add_run(f'  [{count}]')

# Сохранить измененный документ
doc_text.save('text_with_ref.docx')